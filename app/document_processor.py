import os
import os
import json
import random
import re
import tempfile
import pandas as pd
from werkzeug.utils import secure_filename

# Проверяем наличие необходимых библиотек
try:
    import PyPDF2
    import docx
    import pandas as pd
    PDF_DOCX_SUPPORT = True
except ImportError:
    PDF_DOCX_SUPPORT = False
    print("WARNING: PyPDF2, python-docx, or pandas not installed. PDF/DOCX processing will be disabled.")

def extract_text_from_pdf(file_path):
    """Извлекает текст из PDF файла с улучшенной обработкой таблиц."""
    if not PDF_DOCX_SUPPORT:
        return "PDF обработка не поддерживается. Установите PyPDF2."
    
    text = ""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                
                # Преобразуем текст в табличный формат
                lines = page_text.split('\n')
                processed_lines = []
                
                for line in lines:
                    # Ищем строки, которые могут быть частью таблицы
                    # Заменяем множественные пробелы на табуляцию для лучшего разделения колонок
                    processed_line = re.sub(r'\s{2,}', '\t', line)
                    processed_lines.append(processed_line)
                
                text += '\n'.join(processed_lines) + '\n'
    except Exception as e:
        text = f"Ошибка при чтении PDF: {str(e)}"
    
    return text

def extract_text_from_docx(file_path):
    """Извлекает текст из DOCX файла."""
    if not PDF_DOCX_SUPPORT:
        return "DOCX обработка не поддерживается. Установите python-docx."
    
    text = ""
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        
        # Извлекаем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " | "
                text += "\n"
    except Exception as e:
        text = f"Ошибка при чтении DOCX: {str(e)}"
    
    return text

def extract_glossary_items(text):
    """
    Извлекает словарные элементы из текста.
    Предполагается формат таблицы с узбекскими, русскими и английскими терминами.
    """
    items = []
    
    # Разбиваем текст на строки
    lines = text.split('\n')
    
    # Сначала пробуем обработать табличные данные с табуляцией
    for line in lines:
        # Ищем строки с табуляцией
        if '\t' in line:
            parts = line.split('\t')
            if len(parts) >= 3:  # Нам нужно минимум 3 колонки
                # Предполагаем, что первая колонка - узбекский, вторая - русский, третья - английский
                term = parts[0].strip()
                russian = parts[1].strip() if len(parts) > 1 else ""
                english = parts[2].strip() if len(parts) > 2 else ""
                
                # Пропускаем заголовки и пустые строки
                if term and term.lower() not in ['term', 'термин', '№', 'o\'zbek', 'озбек', ''] and \
                   (russian or english):
                    # Если термин начинается с цифры и тире, удаляем эту часть
                    term = re.sub(r'^\d+[-\s]*', '', term)
                    
                    items.append({
                        'term': term,
                        'russian_translation': russian,
                        'english_translation': english
                    })
    
    # Если не нашли табличные данные с табуляцией, пробуем другие методы
    if not items:
        # Пробуем найти таблицу с разделителями |
        table_pattern = re.compile(r'([^|\n]+)\|([^|\n]+)\|([^|\n]+)')
        table_matches = table_pattern.findall(text)
        
        if table_matches:
            for match in table_matches:
                term = match[0].strip()
                russian = match[1].strip()
                english = match[2].strip()
                
                # Пропускаем заголовки и пустые строки
                if term and term.lower() not in ['term', 'термин', '№', 'o\'zbek', 'озбек', ''] and \
                   (russian or english):
                    # Если термин начинается с цифры и тире, удаляем эту часть
                    term = re.sub(r'^\d+[-\s]*', '', term)
                    
                    items.append({
                        'term': term,
                        'russian_translation': russian,
                        'english_translation': english
                    })
    
    # Если все еще нет элементов, пробуем разбить по пробелам
    if not items:
        # Ищем строки, которые могут быть частью таблицы
        for line in lines:
            # Разделяем строку по множественным пробелам
            parts = re.split(r'\s{2,}', line)
            if len(parts) >= 3:  # Нам нужно минимум 3 колонки
                term = parts[0].strip()
                russian = parts[1].strip()
                english = parts[2].strip()
                
                # Пропускаем заголовки и пустые строки
                if term and term.lower() not in ['term', 'термин', '№', 'o\'zbek', 'озбек', ''] and \
                   (russian or english):
                    # Если термин начинается с цифры и тире, удаляем эту часть
                    term = re.sub(r'^\d+[-\s]*', '', term)
                    
                    items.append({
                        'term': term,
                        'russian_translation': russian,
                        'english_translation': english
                    })
    
    # Если все еще нет элементов, пробуем искать по шаблону
    if not items:
        # Специальный шаблон для таблицы с изображения
        pattern = r'(\w+[\w\s-]*?)\s+([А-Яа-я\w\s-]+)\s+([А-Яа-яA-Za-z\w\s-]+)'
        matches = re.findall(pattern, text)
        
        for match in matches:
            term = match[0].strip()
            russian = match[1].strip()
            english = match[2].strip()
            
            # Пропускаем заголовки и пустые строки
            if term and term.lower() not in ['term', 'термин', '№', 'o\'zbek', 'озбек', ''] and \
               (russian or english):
                # Если термин начинается с цифры и тире, удаляем эту часть
                term = re.sub(r'^\d+[-\s]*', '', term)
                
                items.append({
                    'term': term,
                    'russian_translation': russian,
                    'english_translation': english
                })
    
    # Удаляем дубликаты
    unique_items = []
    seen_terms = set()
    
    for item in items:
        if item['term'].lower() not in seen_terms:
            seen_terms.add(item['term'].lower())
            unique_items.append(item)
    
    return unique_items

def create_glossary_test_questions(glossary_items, language='russian'):
    """
    Создает разнообразные вопросы для теста на основе словарных элементов.
    Создает отдельные тесты для русского и английского языков.
    
    Args:
        glossary_items: Список элементов словаря
        language: Язык теста ('russian' или 'english')
    """
    questions = []
    
    # Фильтруем элементы словаря, удаляя неподходящие
    valid_items = []
    for item in glossary_items:
        # Проверяем, что термин и переводы не пустые и не числа
        if item['term'] and not item['term'].isdigit():
            if language == 'russian' and item['russian_translation'] and not item['russian_translation'].isdigit():
                valid_items.append(item)
            elif language == 'english' and item['english_translation'] and not item['english_translation'].isdigit():
                valid_items.append(item)
    
    if not valid_items or len(valid_items) < 4:  # Нужно минимум 4 элемента для создания тестов с вариантами
        return questions
    
    # Создаем вопросы только с узбекского на русский/английский
    for item in valid_items:
        if language == 'russian':
            # Тип 1: Перевод с узбекского на русский
            if item['russian_translation'] and len(item['russian_translation']) > 1:
                # Фильтруем только те элементы, у которых есть русский перевод
                valid_russian_items = [i for i in valid_items if i['term'] != item['term'] and 
                                      i['russian_translation'] and len(i['russian_translation']) > 1]
                
                if len(valid_russian_items) >= 3:
                    # Выбираем неправильные варианты, только на русском языке
                    # Проверяем, что все ответы на русском языке
                    wrong_answers = [i['russian_translation'] for i in random.sample(valid_russian_items, 3)]
                    
                    # Проверяем, что все варианты на русском языке
                    # Используем кириллицу для проверки русского языка
                    cyrillic_pattern = re.compile('[Ѐ-ӿ]')
                    latin_pattern = re.compile('[a-zA-Z]')
                    
                    # Фильтруем варианты, чтобы оставить только те, что содержат кириллицу
                    # и не содержат много латинских букв
                    wrong_answers = [ans for ans in wrong_answers if cyrillic_pattern.search(ans) and 
                                    (not latin_pattern.search(ans) or 
                                     len(re.findall(latin_pattern, ans)) < len(re.findall(cyrillic_pattern, ans)) / 3)]
                    
                    # Если недостаточно вариантов, добавляем стандартные русские слова
                    default_russian_options = ["Одежда", "Ткань", "Фасон", "Стиль", "Материал", "Деталь", "Швейный термин"]
                    while len(wrong_answers) < 3:
                        random_option = random.choice(default_russian_options)
                        if random_option not in wrong_answers and random_option != item['russian_translation']:
                            wrong_answers.append(random_option)
                    
                    options = wrong_answers + [item['russian_translation']]
                    random.shuffle(options)
                    
                    question = {
                        'text': f"Quyidagi atama uchun rus tilidagi tarjimani tanlang: {item['term']}",
                        'type': 'single_choice',
                        'options': {f"option_{i+1}": option for i, option in enumerate(options)},
                        'correct_answer': next(key for key, value in {f"option_{i+1}": option for i, option in enumerate(options)}.items() 
                                              if value == item['russian_translation'])
                    }
                    questions.append(question)
        
        elif language == 'english':
            # Тип 1: Перевод с узбекского на английский
            if item['english_translation'] and len(item['english_translation']) > 1:
                # Фильтруем только те элементы, у которых есть английский перевод
                valid_english_items = [i for i in valid_items if i['term'] != item['term'] and 
                                      i['english_translation'] and len(i['english_translation']) > 1]
            
                if len(valid_english_items) >= 3:
                    # Выбираем неправильные варианты, только на английском языке
                    wrong_answers_eng = [i['english_translation'] for i in random.sample(valid_english_items, 3)]
                    
                    # Проверяем, что все варианты на английском языке
                    # Используем латиницу для проверки английского языка и исключаем кириллицу
                    cyrillic_pattern = re.compile('[Ѐ-ӿ]')
                    latin_pattern = re.compile('[a-zA-Z]')
                    
                    # Фильтруем варианты, чтобы оставить только те, что содержат латинские буквы
                    # и не содержат кириллицу
                    wrong_answers_eng = [ans for ans in wrong_answers_eng if latin_pattern.search(ans) and not cyrillic_pattern.search(ans)]
                    
                    # Если недостаточно вариантов, добавляем стандартные английские слова
                    default_english_options = ["Clothing", "Fabric", "Style", "Fashion", "Material", "Garment", "Sewing term"]
                    while len(wrong_answers_eng) < 3:
                        random_option = random.choice(default_english_options)
                        if random_option not in wrong_answers_eng and random_option != item['english_translation']:
                            wrong_answers_eng.append(random_option)
                    
                    options_eng = wrong_answers_eng + [item['english_translation']]
                    random.shuffle(options_eng)
                    
                    question_eng = {
                        'text': f"Quyidagi atama uchun ingliz tilidagi tarjimani tanlang: {item['term']}",
                        'type': 'single_choice',
                        'options': {f"option_{i+1}": option for i, option in enumerate(options_eng)},
                        'correct_answer': next(key for key, value in {f"option_{i+1}": option for i, option in enumerate(options_eng)}.items() 
                                              if value == item['english_translation'])
                    }
                    questions.append(question_eng)
            
    # Добавляем дополнительные вопросы для разнообразия
    
    # Добавляем вопросы с вводом текста (если есть достаточно элементов)
    if language == 'russian':
        # Для русского теста - вопросы с узбекского на русский
        text_input_items = [item for item in valid_items if item['russian_translation'] and len(item['russian_translation']) > 1 and 
                            not any(c.isdigit() for c in item['term']) and len(item['term']) > 2]
        
        # Проверяем, что ответы на русском языке
        cyrillic_pattern = re.compile('[Ѐ-ӿ]')
        latin_pattern = re.compile('[a-zA-Z]')
        
        # Фильтруем элементы, чтобы оставить только те, что содержат кириллицу
        # и не содержат много латинских букв
        text_input_items = [item for item in text_input_items if cyrillic_pattern.search(item['russian_translation']) and 
                           (not latin_pattern.search(item['russian_translation']) or 
                            len(re.findall(latin_pattern, item['russian_translation'])) < 
                            len(re.findall(cyrillic_pattern, item['russian_translation'])) / 3)]
        
        if len(text_input_items) >= 3:
            # Выбираем несколько случайных элементов для вопросов с вводом текста
            for item in random.sample(text_input_items, min(3, len(text_input_items))):
                # Вопрос с вводом текста - перевод с узбекского на русский
                question_text_input = {
                    'text': f"Quyidagi atama uchun rus tilidagi tarjimani kiriting: {item['term']}",
                    'type': 'text_input',
                    'options': {},  # Пустой словарь для совместимости с существующим кодом
                    'correct_answer': item['russian_translation']
                }
                questions.append(question_text_input)
    
    elif language == 'english':
        # Для английского теста - вопросы с узбекского на английский
        text_input_items = [item for item in valid_items if item['english_translation'] and len(item['english_translation']) > 1 and 
                            not any(c.isdigit() for c in item['term']) and len(item['term']) > 2]
        
        # Проверяем, что ответы на английском языке
        cyrillic_pattern = re.compile('[Ѐ-ӿ]')
        latin_pattern = re.compile('[a-zA-Z]')
        
        # Фильтруем элементы, чтобы оставить только те, что содержат латинские буквы
        # и не содержат кириллицу
        text_input_items = [item for item in text_input_items if latin_pattern.search(item['english_translation']) and 
                           not cyrillic_pattern.search(item['english_translation'])]
        
        if len(text_input_items) >= 3:
            # Выбираем несколько случайных элементов для вопросов с вводом текста
            for item in random.sample(text_input_items, min(3, len(text_input_items))):
                # Вопрос с вводом текста - перевод с узбекского на английский
                question_text_input = {
                    'text': f"Quyidagi atama uchun ingliz tilidagi tarjimani kiriting: {item['term']}",
                    'type': 'text_input',
                    'options': {},
                    'correct_answer': item['english_translation']
                }
                questions.append(question_text_input)
    
    # Перемешиваем вопросы и ограничиваем их количество
    random.shuffle(questions)
    
    # Ограничиваем количество вопросов до разумного количества (20-30 вопросов максимум)
    max_questions = min(len(questions), 30)
    return questions[:max_questions]

def process_glossary_file(file, upload_folder):
    """
    Обрабатывает загруженный файл словаря и извлекает словарные элементы.
    """
    if not PDF_DOCX_SUPPORT:
        return [], "Обработка документов не поддерживается. Для работы с PDF и DOCX файлами установите библиотеки PyPDF2, python-docx и pandas. Вы можете добавлять термины вручную через форму на странице."
    
    # Сохраняем файл во временную директорию
    filename = secure_filename(file.filename)
    temp_dir = tempfile.mkdtemp()
    file_path = os.path.join(temp_dir, filename)
    file.save(file_path)
    
    # Определяем тип файла и извлекаем текст
    if filename.lower().endswith('.pdf'):
        text = extract_text_from_pdf(file_path)
    elif filename.lower().endswith('.docx'):
        text = extract_text_from_docx(file_path)
    else:
        os.remove(file_path)
        os.rmdir(temp_dir)
        return [], "Неподдерживаемый формат файла. Поддерживаются только PDF и DOCX."
    
    # Извлекаем словарные элементы
    items = extract_glossary_items(text)
    
    # Проверяем и исправляем переводы
    filtered_items = []
    for item in items:
        # Фильтруем заголовки столбцов и другие некорректные записи
        if 'term' in item and item['term']:
            # Пропускаем записи, которые похожи на заголовки столбцов
            if item['term'].lower() in ['o\'zbek', '№ o\'zbek', 'uzbek', 'term', 'термин']:
                continue
            
            # Пропускаем записи, где перевод похож на заголовок столбца
            if 'russian_translation' in item and item['russian_translation']:
                if item['russian_translation'].lower() in ['rus', 'rus ingliz', 'русский', 'russian', 'translation']:
                    continue
            
            if 'english_translation' in item and item['english_translation']:
                if item['english_translation'].lower() in ['ingliz', 'english', 'translation']:
                    continue
            
            # Проверяем, что русский перевод не содержит английских слов
            if 'russian_translation' in item and item['russian_translation']:
                # Если в русском переводе есть английские слова и нет русских букв, то это вероятно английский перевод
                if all(ord(c) < 1072 or ord(c) > 1103 for c in item['russian_translation'].lower() if c.isalpha()):
                    # Если английский перевод пуст, перемещаем русский в английский
                    if not item.get('english_translation'):
                        item['english_translation'] = item['russian_translation']
                        item['russian_translation'] = ""
            
            # Проверяем специальные случаи переводов
            if 'russian_translation' in item and item['russian_translation'] == "Constructive line":
                item['russian_translation'] = "Конструктивная линия"
                if not item.get('english_translation'):
                    item['english_translation'] = "Constructive line"
            
            # Убедиться, что английский перевод не содержит узбекских или русских слов
            if 'english_translation' in item and item['english_translation']:
                # Если в английском переводе есть русские буквы, то это вероятно русский перевод
                if any(1072 <= ord(c) <= 1103 for c in item['english_translation'].lower() if c.isalpha()):
                    if not item.get('russian_translation'):
                        item['russian_translation'] = item['english_translation']
                        item['english_translation'] = ""
            
            # Добавляем только корректные записи
            filtered_items.append(item)
    
    # Удаляем временный файл
    os.remove(file_path)
    os.rmdir(temp_dir)
    
    if not filtered_items:
        return [], "Не удалось извлечь словарные элементы из файла. Убедитесь, что файл содержит таблицу с терминами и переводами."
    
    return filtered_items, None
